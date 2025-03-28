from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
import re

def set_run_font(run, font_name, font_size):
    """设置单个 run 的字体名称、字号，并指定中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def update_paragraph_font(paragraph, font_name, font_size):
    """遍历段落所有 run 统一设置字体"""
    for run in paragraph.runs:
        set_run_font(run, font_name, font_size)

# 定义日志输入文件
report_file = "mysql_system_report.log"

# 读取整个日志文件
with open(report_file, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# ----------------------------
# 提取巡检日期
# 日志第一行格式应为 "MySQL & 操作系统巡检报告 - <日期字符串>"
if lines:
    header_line = lines[0].strip()
    if " - " in header_line:
        date_str_raw = header_line.split(" - ", 1)[1]
        try:
            # 假设日期格式为 "%a %b %d %H:%M:%S %Y"
            inspection_date = datetime.strptime(date_str_raw, "%a %b %d %H:%M:%S %Y")
            date_filename = inspection_date.strftime("%Y%m%d_%H%M%S")
        except Exception:
            date_filename = datetime.now().strftime("%Y%m%d_%H%M%S")
    else:
        date_filename = datetime.now().strftime("%Y%m%d_%H%M%S")
else:
    date_filename = datetime.now().strftime("%Y%m%d_%H%M%S")

# ----------------------------
# 提取主机名
# 查找包含 "主机名：" 的行（忽略前面的"##"或其它标识）
hostname = "unknown"
for line in lines:
    if "主机名：" in line:
        # 分割时以"主机名："为标识，取后面的部分
        hostname = line.split("主机名：", 1)[1].strip()
        break
# 清洗主机名（去除斜杠、空格等非法字符）
hostname = hostname.replace("/", "-").replace(" ", "_")

# ----------------------------
# 提取 IP 地址
# 查找包含 "IP 地址及网络接口：" 的行，后续行中找到首个以 "inet" 开头且非127.* 的IP地址
ip_address = "unknown"
found_section = False
for line in lines:
    if "IP 地址及网络接口：" in line:
        found_section = True
        continue
    if found_section:
        parts = line.split()
        if parts and parts[0] == "inet":
            candidate = parts[1]
            if "/" in candidate:
                candidate = candidate.split("/")[0]
            if not candidate.startswith("127."):
                ip_address = candidate
                break
# 清洗 IP 地址
ip_address = ip_address.replace("/", "-").replace(" ", "_")

# 构造 Word 文件名，格式为 "<主机名>_<IP地址>_<YYYYMMDD_HHMMSS>.docx"
generated_word_file = f"{hostname}_{ip_address}_{date_filename}.docx"

# ----------------------------
# 创建 Word 文档
doc = Document()

# 添加文档标题（一级标题）
title = doc.add_heading('MySQL & 操作系统巡检报告', level=1)
update_paragraph_font(title, "新宋体", 16)

p_time = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
update_paragraph_font(p_time, "新宋体", 12)
p_sep = doc.add_paragraph("=====================================")
update_paragraph_font(p_sep, "新宋体", 12)

# 判断二级标题采用简单字符串方法：若行以 "##" 开头，则视为二级标题
for line in lines:
    line = line.strip()
    # 跳过不需要导入的提示部分
    if line.startswith("#【4】") or "本脚本不自动调用 Python 生成报告" in line:
        continue
    if line.startswith("#【"):
        p = doc.add_paragraph(line, style='Heading 1')
        update_paragraph_font(p, "新宋体", 16)
    elif line.startswith("##"):
        # 去除前缀 "##" 后作为二级标题
        new_line = line[2:].strip()
        p = doc.add_paragraph(new_line, style='Heading 2')
        update_paragraph_font(p, "新宋体", 14)
    elif "ERROR" in line or "WARNING" in line:
        p = doc.add_paragraph(line, style='Intense Quote')
        update_paragraph_font(p, "新宋体", 12)
    else:
        p = doc.add_paragraph(line)
        update_paragraph_font(p, "新宋体", 12)

doc.save(generated_word_file)
print(f"巡检报告已生成：{generated_word_file}")
